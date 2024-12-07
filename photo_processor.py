import sys
import os

from PyQt6.QtWidgets import QMainWindow, QFileDialog, QMessageBox
from PyQt6.QtGui import QPixmap, QIntValidator
from PyQt6.QtCore import Qt, QObject, QEvent
from PyQt6 import uic

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx import Document

desktop = os.path.join(os.path.expanduser("~"))

base_path = getattr(sys, '_MEIPASS', os.path.abspath("."))
ui_path = os.path.join(base_path, 'photo_processor.ui')
photo_processor_ui, classinfo = uic.loadUiType(ui_path)

class img_path_drop_handler(QObject):
    def eventFilter(self, watched, event):
        if event.type() == QEvent.Type.DragEnter:
            event.accept()
        if event.type() == QEvent.Type.Drop:
            md = event.mimeData()
            for url in event.mimeData().urls():
                url = url
            if md.hasUrls():
                watched.setText(url.toLocalFile())
                return True
        return super().eventFilter(watched, event)

class PhotoID_Processor(QMainWindow, photo_processor_ui):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.setFixedSize(self.size())
    
        self.document = Document()

        self.selected_file = None

        self.set_img_path()

        self.id_type.addItems(['','2x2', '1x1'])

        validator = QIntValidator(1, 99)
        self.num_of_photos.setValidator(validator)

        self.browse_button.clicked.connect(self.open_file)
        self.process_button.clicked.connect(self.handle_file)
        
        self.num_of_photos.textChanged.connect(self.validate_input)

        self.img_path.installEventFilter(img_path_drop_handler(self))
        self.img_path.textChanged.connect(lambda: self.open_file('img_path'))

    def set_img_path(self):
        if not self.selected_file:
            self.img_path.setText(desktop)
        else:
            self.img_path.setText(self.selected_file)

    def open_file(self, field):
        if field == 'img_path':
            self.selected_file = self.img_path.text()
        else:
            self.selected_file = self.open_dialog('open_file')
        self.preview_image()
        self.set_img_path()
    
    def validate_input(self):
        if self.num_of_photos.text() == '0':
            self.num_of_photos.clear()

    def handle_file(self):
        id_type = self.id_type.currentText()
        num_of_photos = self.num_of_photos.text()

        if not self.selected_file or self.selected_file == desktop:
            QMessageBox.warning(self, 'Warning', 'Please select an image!')
            return
        if not num_of_photos:
            QMessageBox.warning(self, 'Warning', 'Please input how many photos to process!')
            return

        else:
            match id_type:
                case '2x2':                    
                    self.create_photoID(2, num_of_photos, self.selected_file, id_type)

                case '1x1':
                    self.create_photoID(1, num_of_photos, self.selected_file, id_type)
                
                case '':
                    QMessageBox.warning(self, 'Warning', 'Please select an ID TYPE!')

    def create_photoID(self, size, num_of_photos, selected_file, id_type):
        try:
            photos = int(num_of_photos)
            counter = 0
            
            p = self.document.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = p.add_run()
            
            for id in range(photos):
                counter += 1
                
                r.add_picture(selected_file, width=Inches(size), height=Inches(size))
                match (counter, id_type):
                    case (2, '2x2'):
                        r.add_break()
                        counter = 0
                    case (4, '1x1'):
                        r.add_break()
                        counter = 0
        except Exception as e:
            return

        self.save_file()

    def save_file(self):
        try:
            save_to = self.open_dialog('save_file')
            self.document.save(save_to)
        except Exception as e:
            QMessageBox.warning(self, 'Warning', 'Word document is currently open in another tab, please close!')
            return
    
    def preview_image(self):
        pixmap = QPixmap(self.selected_file)
        
        if pixmap.isNull():
            return
        
        scaled_pixmap = pixmap.scaled(self.preview_img.size(), Qt.AspectRatioMode.KeepAspectRatio)
        self.preview_img.setPixmap(scaled_pixmap)
        
    def open_dialog(self, mode):
        dialog = QFileDialog(self)

        match mode:
            case 'open_file':
                filename, _ = dialog.getOpenFileName(
                    self,
                    "Open File",
                    desktop,
                    "Images (*.png *.jpg)"
                )
                return filename
                          
            case 'save_file':
                filename, _ = dialog.getSaveFileName(
                    self,
                    "Save File",
                    desktop,
                    "Word Documents (*.docx)"
                )
                return filename